"""
Job Processor - Optimized for Ollama Translation
=================================================
- Uses Ollama with qwen3:8b as primary FREE translator
- Auto-detects GPU for acceleration
- Smart chunking for structured vs narrative pages
- Tracks provider and speed info
"""

import asyncio
import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Tuple

import structlog

from app.config import OCRProvider, TranslationProvider, get_settings
from app.models import JobStatus
from app.services.ocr_service import OCRService, OCRProgress, ImageExtractor
from app.services.translation_service import TranslationService

logger = structlog.get_logger()


def check_gpu_available() -> bool:
    """Check if GPU is available for acceleration."""
    try:
        import torch
        return torch.cuda.is_available()
    except ImportError:
        return False


class JobProcessor:
    """Processes translation jobs with Ollama (qwen3:8b or configured model)."""

    # Chunk size for page-level splitting (Ollama will further split for 4K context)
    # ~3000 chars = roughly 1 page, good for progress tracking
    OPTIMIZED_CHUNK_SIZE = 3000
    
    def __init__(self, db_session=None):
        self.settings = get_settings()
        self.ocr_service = OCRService()
        self.translation_service = TranslationService()
        self.db = db_session
        self.gpu_available = check_gpu_available()
        
        if self.gpu_available:
            logger.info("GPU detected - NLLB will use CUDA acceleration")
        else:
            logger.warning("No GPU detected - NLLB will run on CPU (slower)")
    
    

    def _debug_dir(self, job_id: str) -> Path:
        settings = get_settings()
        base = Path(settings.debug_output_dir)
        d = base / str(job_id)
        d.mkdir(parents=True, exist_ok=True)
        return d

    def _debug_write_text(self, job_id: str, filename: str, content: str):
        settings = get_settings()
        if not getattr(settings, "debug_translation", False):
            return
        d = self._debug_dir(job_id)
        max_chars = int(getattr(settings, "debug_max_chars_per_page", 200000))
        safe = content if content is not None else ""
        if max_chars > 0 and len(safe) > max_chars:
            safe = safe[:max_chars] + "\n\n[TRUNCATED]\n"
        (d / filename).write_text(safe, encoding="utf-8", errors="ignore")

    def _debug_write_json(self, job_id: str, filename: str, data: dict):
        settings = get_settings()
        if not getattr(settings, "debug_translation", False):
            return
        d = self._debug_dir(job_id)
        (d / filename).write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8", errors="ignore")
    async def process_job(
        self,
        job_id: str,
        input_path: str,
        source_lang: Optional[str],
        target_lang: str,
        ocr_provider: Optional[str] = None,
        translation_provider: Optional[str] = None,
        status_callback=None
    ) -> dict:
        """
        Process translation job with speed optimization.
        """
        try:
            # Determine provider info for user feedback
            provider_name = translation_provider or "ollama"  # Default to Ollama now

            # Check if Ollama is available
            ollama_available = False
            ollama_model = self.settings.ollama_model  # Get configured model name
            try:
                ollama_translator = self.translation_service.providers.get(TranslationProvider.OLLAMA)
                if ollama_translator:
                    ollama_available = ollama_translator.is_available()
                    logger.info(f"Ollama availability check: {ollama_available}")
                else:
                    logger.warning("OllamaTranslator not found in providers")
            except Exception as e:
                logger.error(f"Error checking Ollama availability: {e}")

            # Auto-select provider based on GPU availability
            # Priority: GPU + Ollama > GPU + NLLB > DeepSeek API > CPU NLLB (slow)
            if not translation_provider:
                if self.gpu_available:
                    # GPU detected - free local models will be fast
                    if ollama_available:
                        provider_name = "ollama"
                        logger.info(f"GPU detected! Using FREE Ollama with {ollama_model} (fast, high accuracy)")
                    else:
                        provider_name = "nllb"
                        logger.info("GPU detected! Using NLLB with CUDA acceleration (fast)")
                else:
                    # No GPU - warn about slow performance, recommend API
                    if ollama_available:
                        provider_name = "ollama"
                        logger.warning(f"No GPU detected. Ollama/{ollama_model} will run on CPU (slower). Consider using DeepSeek API for faster results.")
                    else:
                        provider_name = "nllb"
                        logger.warning("No GPU detected. NLLB will run on CPU (SLOW). Recommend using DeepSeek API (set DEEPSEEK_API_KEY) for faster translation.")

            # Configure device info based on provider
            if provider_name == "ollama":
                device_info = f"{ollama_model} ({'GPU' if self.gpu_available else 'CPU'})"
                logger.info(f"Translation provider: Ollama with {ollama_model} on {'GPU' if self.gpu_available else 'CPU'}")
            elif provider_name == "nllb":
                # If using NLLB and GPU available, configure for CUDA
                if self.gpu_available:
                    try:
                        self.settings.nllb_device = "cuda"
                    except Exception:
                        pass
                device_info = "GPU (CUDA)" if self.gpu_available else "CPU (slow)"
                logger.info(f"Using NLLB on {device_info}")
            else:
                device_info = "API"
            
            # ===== STEP 1: EXTRACT TEXT (0-20%) =====
            if status_callback:
                await status_callback(job_id, JobStatus.EXTRACTING, 0, "Analyzing document...")
            
            loop = asyncio.get_event_loop()
            
            def ocr_progress_wrapper(progress: OCRProgress):
                if status_callback:
                    pct = (progress.current_page / max(progress.total_pages, 1)) * 20
                    asyncio.run_coroutine_threadsafe(
                        status_callback(job_id, JobStatus.EXTRACTING, pct,
                            f"Extracting page {progress.current_page}/{progress.total_pages}"),
                        loop
                    )
            
            ocr_prov = OCRProvider(ocr_provider) if ocr_provider else None
            
            extracted_text, page_count, ocr_used, page_texts = await self.ocr_service.extract_text(
                input_path, ocr_prov, progress_callback=ocr_progress_wrapper
            )
            # Debug instrumentation: save extracted text + per-page language detection
            settings = get_settings()
            if getattr(settings, "debug_translation", False):
                self._debug_write_text(job_id, "extracted_full.txt", extracted_text or "")
                self._debug_write_json(job_id, "extraction_summary.json", {
                    "input_path": str(input_path),
                    "page_count": page_count,
                    "page_texts_len": len(page_texts) if page_texts is not None else None,
                    "ocr_used": ocr_used,
                    "ocr_provider_requested": str(ocr_provider) if ocr_provider else None,
                    "translation_provider_requested": str(translation_provider) if translation_provider else None,
                    "device_info": device_info,
                })
                for i, page_text in enumerate(page_texts or []):
                    self._debug_write_text(job_id, f"page_{i:03d}_extracted.txt", page_text or "")
                    lang_details = await self.translation_service.detect_language_details(page_text or "")
                    self._debug_write_json(job_id, f"page_{i:03d}_lang.json", lang_details)

            
            if not extracted_text.strip():
                raise ValueError("No text could be extracted from the document")
            
            logger.info(f"Extraction complete: {page_count} pages, {len(extracted_text)} chars, method={ocr_used}")
            
            # ===== STEP 1B: EXTRACT IMAGES (NEW) =====
            if status_callback:
                await status_callback(job_id, JobStatus.EXTRACTING, 18, "Extracting images...")
            
            page_images = []
            try:
                page_images = await ImageExtractor.extract_images_from_pdf(input_path)
                total_images = sum(len(imgs) for imgs in page_images)
                logger.info(f"Image extraction complete: {total_images} images from {len(page_images)} pages")
            except Exception as e:
                logger.warning(f"Image extraction failed (continuing without images): {e}")
                page_images = [[] for _ in range(page_count)]  # Empty lists as fallback
            
            # ===== STEP 2: DETECT LANGUAGE (20-25%) =====
            if status_callback:
                await status_callback(job_id, JobStatus.EXTRACTING, 22, "Detecting language...")
            
            if not source_lang:
                source_lang = await self.translation_service.detect_language(extracted_text)
                logger.info(f"Detected language: {source_lang}")
            
            # ===== STEP 3: BATCH TRANSLATION (25-90%) =====
            speed_note = f"Using {provider_name} on {device_info}"
            if status_callback:
                await status_callback(job_id, JobStatus.TRANSLATING, 25, f"Preparing translation... ({speed_note})")
            
            trans_prov = TranslationProvider(translation_provider) if translation_provider else None
            
            # Separate pages: needs translation vs skip
            pages_to_translate = []
            pages_to_skip = []
            
            for i, page_text in enumerate(page_texts):
                if not page_text.strip():
                    pages_to_skip.append((i, ""))
                elif not self.translation_service.needs_translation(page_text, target_lang):
                    pages_to_skip.append((i, page_text))
                else:
                    pages_to_translate.append((i, page_text))
            
            pages_skipped = len(pages_to_skip)
            total_to_translate = len(pages_to_translate)
            
            logger.info(f"Speed optimization: {total_to_translate} pages to translate, {pages_skipped} skipped")
            
            if status_callback:
                await status_callback(
                    job_id, JobStatus.TRANSLATING, 28,
                    f"Translating {total_to_translate} pages ({pages_skipped} already English)...",
                    {"pages_skipped": pages_skipped, "total_pages": page_count}
                )
            
            # Initialize results
            translated_results = {idx: text for idx, text in pages_to_skip}
            
            # Translate pages
            for i, (idx, page_text) in enumerate(pages_to_translate):
                translated_text = await self._translate_page_optimized(
                    page_text, source_lang, target_lang, trans_prov
                )
                translated_results[idx] = translated_text
                
                # Progress update
                progress = 28 + (62 * (i + 1) / max(total_to_translate, 1))
                
                if status_callback:
                    await status_callback(
                        job_id, JobStatus.TRANSLATING, min(progress, 90),
                        f"Translated {i + 1}/{total_to_translate} pages... ({speed_note})",
                        {"pages_translated": i + 1}
                    )
            
            # Reconstruct in order
            translated_pages = [translated_results.get(i, "") for i in range(page_count)]
            # Debug instrumentation: save translated pages
            settings = get_settings()
            if getattr(settings, "debug_translation", False):
                for i, t in enumerate(translated_pages or []):
                    self._debug_write_text(job_id, f"page_{i:03d}_translated.txt", t or "")
                self._debug_write_text(job_id, "translated_full.txt", "\n\n".join(translated_pages or []))

            
            logger.info(f"Translation complete: {total_to_translate} translated, {pages_skipped} skipped")
            
            # ===== STEP 4: CREATE WORD DOCUMENT (90-100%) =====
            if status_callback:
                await status_callback(job_id, JobStatus.TRANSLATING, 92, "Creating Word document...")
            
            output_filename = self._generate_output_filename(input_path, target_lang)
            output_path = Path(self.settings.output_dir) / output_filename
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            self._create_word_document(
                translated_pages, 
                str(output_path), 
                page_count, 
                source_lang, 
                target_lang,
                total_to_translate,
                pages_skipped,
                page_images  # NEW: Pass images
            )
            
            # ===== DONE =====
            if status_callback:
                await status_callback(
                    job_id, JobStatus.COMPLETED, 100, "Translation complete!",
                    {
                        "output_path": str(output_path),
                        "total_pages": page_count,
                        "pages_translated": total_to_translate,
                        "pages_skipped": pages_skipped
                    }
                )
            
            return {
                "success": True,
                "output_path": str(output_path),
                "page_count": page_count,
                "pages_translated": total_to_translate,
                "pages_skipped": pages_skipped,
                "source_language": source_lang,
                "target_language": target_lang,
                "ocr_provider": ocr_used,
                "translation_provider": provider_name,
                "gpu_used": self.gpu_available and provider_name == "nllb"
            }
            
        except Exception as e:
            import traceback
            error_msg = str(e)
            logger.error(f"Job {job_id} failed: {error_msg}", traceback=traceback.format_exc())
            if status_callback:
                await status_callback(job_id, JobStatus.FAILED, 0, f"Failed: {error_msg}")
            raise
    
    async def _translate_page_optimized(
        self,
        page_text: str,
        source_lang: str,
        target_lang: str,
        provider: Optional[TranslationProvider]
    ) -> str:
        """
        Translate a page using WHOLE-PAGE translation for speed.

        Layout-preserving extraction already maintains TOC structure,
        so we use whole-page translation for everything (faster).
        Pre-replacements handle critical terms that LLMs often mistranslate.
        """
        # Step 1: Remove duplicate bilingual content (same text in both languages)
        page_text = self._deduplicate_bilingual_content(page_text, source_lang, target_lang)

        # Step 2: Apply critical term pre-replacements for accuracy
        # This handles terms like Riphean, Kovel, Fore-Dobruja that LLMs often get wrong
        page_text = self._apply_critical_term_replacements(page_text)

        # Step 3: ALWAYS use whole-page translation (fastest)
        # Layout-preserving extraction already maintains structure for TOC pages
        # Line-by-line is too slow with deepseek-r1 (each line triggers "thinking")
        if self.translation_service.needs_translation(page_text, target_lang):
            result, _, _ = await self.translation_service.translate_if_needed(
                page_text, source_lang, target_lang, provider
            )
            return result
        else:
            return page_text

    def _apply_critical_term_replacements(self, text: str) -> str:
        """
        Pre-replace critical terms that LLMs often mistranslate.
        Applied BEFORE sending to translation to ensure accuracy.
        
        IMPORTANT: Order matters! Full abbreviation lines must be replaced FIRST,
        before individual terms like 'прогин' → 'Trough' can break the match.
        """
        if not text:
            return text

        result = text
        
        # ===== PHASE 1: FULL ABBREVIATION LINES (must be first!) =====
        # These contain terms like 'прогин' that would be replaced in Phase 2
        abbreviation_lines = {
            # Header
            'ПРИЙНЯТІ СКОРОЧЕННЯ': 'ACCEPTED ABBREVIATIONS',
            'Прийняті скорочення': 'Accepted Abbreviations',
            
            # Full abbreviation lines (match the exact format from the document)
            'ВКМ        — Воронезький кристалічний масив': 'VKM — Voronezh Crystalline Massif',
            'ВКМ — Воронезький кристалічний масив': 'VKM — Voronezh Crystalline Massif',
            'ВПМ       — Волино-Подільська монокліналь': 'VPM — Volyn-Podillia Monocline',
            'ВПМ — Волино-Подільська монокліналь': 'VPM — Volyn-Podillia Monocline',
            'Геолком  — Геологічний Комітет': 'Geolcom — Geological Committee',
            'Геолком — Геологічний Комітет': 'Geolcom — Geological Committee',
            'ДДЗ        — Дніпровсько-Донецька западина': 'DDZ — Dnieper-Donets Depression',
            'ДДЗ — Дніпровсько-Донецька западина': 'DDZ — Dnieper-Donets Depression',
            'ДДП        — Доно-Дніпровський прогин': 'DDP — Don-Dnieper Trough',
            'ДДП — Доно-Дніпровський прогин': 'DDP — Don-Dnieper Trough',
            'ДПМ       — Дністровсько-Прутське межиріччя': 'DPM — Dniester-Prut Interfluve',
            'ДПМ — Дністровсько-Прутське межиріччя': 'DPM — Dniester-Prut Interfluve',
            'ЗСШ       — Загальна стратиграфічна шкала': 'GSS — General Stratigraphic Scale',
            'ЗСШ — Загальна стратиграфічна шкала': 'GSS — General Stratigraphic Scale',
            'ІГН          — Інститут геологічних наук': 'IGS — Institute of Geological Sciences',
            'ІГН — Інститут геологічних наук': 'IGS — Institute of Geological Sciences',
            'КСС        — Кореляційна стратиграфічна схема': 'CSS — Correlation Stratigraphic Scheme',
            'КСС — Кореляційна стратиграфічна схема': 'CSS — Correlation Stratigraphic Scheme',
            'ЛБ          — Люблінський басейн': 'LB — Lublin Basin',
            'ЛБ — Люблінський басейн': 'LB — Lublin Basin',
            'ЛПП        — Львівський палеозойський прогин': 'LPT — Lviv Paleozoic Trough',
            'ЛПП — Львівський палеозойський прогин': 'LPT — Lviv Paleozoic Trough',
            'МГК       — Міжнародний геологічний конгрес': 'IGC — International Geological Congress',
            'МГК — Міжнародний геологічний конгрес': 'IGC — International Geological Congress',
            'МСК       — Міжвідомчий стратиграфічний комітет': 'ISC — Interdepartmental Stratigraphic Committee',
            'МСК — Міжвідомчий стратиграфічний комітет': 'ISC — Interdepartmental Stratigraphic Committee',
            'МСКУ      — Міжвідомчий стратиграфічний комітет України': 'ISCU — Interdepartmental Stratigraphic Committee of Ukraine',
            'МСКУ — Міжвідомчий стратиграфічний комітет України': 'ISCU — Interdepartmental Stratigraphic Committee of Ukraine',
            'МСШ      — Міжнародна стратиграфічна шкала': 'ISS — International Stratigraphic Scale',
            'МСШ — Міжнародна стратиграфічна шкала': 'ISS — International Stratigraphic Scale',
            'МФГ       — мікрофауністичний горизонт': 'MFH — Microfaunal Horizon',
            'МФГ — мікрофауністичний горизонт': 'MFH — Microfaunal Horizon',
            'НСКУ      — Національний стратиграфічний комітет України': 'NSCU — National Stratigraphic Committee of Ukraine',
            'НСКУ — Національний стратиграфічний комітет України': 'NSCU — National Stratigraphic Committee of Ukraine',
            'ПКМ       — Приазовський кристалічний масив': 'PKM — Priazovian Crystalline Massif',
            'ПКМ — Приазовський кристалічний масив': 'PKM — Priazovian Crystalline Massif',
            'СЄП       — Східно-Європейська платформа': 'EEP — East European Platform',
            'СЄП — Східно-Європейська платформа': 'EEP — East European Platform',
            'СФЗ       — структурно-фаціальна зона': 'SFZ — Structural-Facies Zone',
            'СФЗ — структурно-фаціальна зона': 'SFZ — Structural-Facies Zone',
            'СФП       — структурно-фаціальна підзона': 'SFS — Structural-Facies Subzone',
            'СФП — структурно-фаціальна підзона': 'SFS — Structural-Facies Subzone',
            'УСС        — Уніфікована схема стратиграфії': 'USS — Unified Stratigraphic Scheme',
            'УСС — Уніфікована схема стратиграфії': 'USS — Unified Stratigraphic Scheme',
            'УЩ         — Український щит': 'US — Ukrainian Shield',
            'УЩ — Український щит': 'US — Ukrainian Shield',
        }
        
        # Apply abbreviation line replacements FIRST (longest matches first)
        for uk_term, en_term in sorted(abbreviation_lines.items(), key=lambda x: -len(x[0])):
            if uk_term in result:
                result = result.replace(uk_term, en_term)
        
        # ===== PHASE 2: INDIVIDUAL TERMS (after full lines are replaced) =====
        individual_terms = {
            # Stratigraphy periods (CRITICAL - often mistranslated)
            'СТРАТИГРАФІЯ РИФЕЮ': 'STRATIGRAPHY OF THE RIPHEAN',
            'Стратиграфія рифею': 'Stratigraphy of the Riphean',
            'РИФЕЙ': 'RIPHEAN',
            'Рифей': 'Riphean',
            'рифею': 'Riphean',
            'рифейськ': 'Riphean',
            # Geographic names
            'Ковельськ': 'Kovel',
            'КОВЕЛЬСЬК': 'KOVEL',
            'Переддобрудзьк': 'Fore-Dobruja',
            'ПЕРЕДДОБРУДЗЬК': 'FORE-DOBRUJA',
            'Переддобружжя': 'Fore-Dobruja',
            # Erathem
            'ЕРАТЕМА': 'ERATHEM',
            'Ератема': 'Erathem',
            'ератема': 'erathem',
            # Vendian
            'ВЕНД': 'VENDIAN',
            'Венд': 'Vendian',
            # Structures (these can break abbreviation lines if applied first!)
            'прогин': 'Trough',
            'ПРОГИН': 'TROUGH',
        }
        
        for uk_term, en_term in individual_terms.items():
            if uk_term in result:
                result = result.replace(uk_term, en_term)

        return result

    @staticmethod
    def _looks_like_toc(text: str) -> bool:
        if not text or len(text) < 200:
            return False
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        if len(lines) < 8:
            return False

        dot_leader_lines = sum(1 for ln in lines if '...' in ln or '…..' in ln or '…' in ln)
        end_number_lines = sum(1 for ln in lines if re.search(r"\b\d{1,4}\s*$", ln))
        short_lines = sum(1 for ln in lines if len(ln) <= 60)

        # Heuristic: TOC pages have many short lines + many trailing numbers and/or dot leaders
        return (end_number_lines >= max(5, int(len(lines) * 0.35))) and (short_lines >= int(len(lines) * 0.6))

    async def _translate_toc_page(
        self,
        page_text: str,
        source_lang: str,
        target_lang: str,
        provider: Optional[TranslationProvider]
    ) -> str:
        """Translate a TOC-like page line-by-line, preserving page numbers."""
        out_lines: list[str] = []
        for raw in page_text.splitlines():
            line = raw.rstrip("\n")
            if not line.strip():
                out_lines.append("")
                continue

            # Preserve trailing page number(s)
            m = re.match(r"^(.*?)(\s+)(\d{1,4})\s*$", line)
            if m:
                left, sep, num = m.group(1), m.group(2), m.group(3)
                if self.translation_service.needs_translation(left, target_lang):
                    translated_left, _, _ = await self.translation_service.translate_if_needed(
                        left, source_lang, target_lang, provider
                    )
                    out_lines.append(f"{translated_left}{sep}{num}")
                else:
                    out_lines.append(line)
                continue

            # Otherwise translate whole line if needed
            if self.translation_service.needs_translation(line, target_lang):
                translated_line, _, _ = await self.translation_service.translate_if_needed(
                    line, source_lang, target_lang, provider
                )
                out_lines.append(translated_line)
            else:
                out_lines.append(line)

        return "\n".join(out_lines)

    def _deduplicate_bilingual_content(self, text: str, source_lang: str, target_lang: str) -> str:
        """
        Remove duplicate content that appears in both source and target language.
        Common in documents with parallel text (e.g., Ukrainian abstract + English abstract).
        """
        if not text or len(text) < 100:
            return text

        paragraphs = re.split(r'\n\s*\n', text)
        if len(paragraphs) < 2:
            return text

        # Check for near-duplicate paragraphs (same semantic content in different languages)
        seen_content = set()
        unique_paragraphs = []

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Normalize for comparison: lowercase, remove punctuation, collapse whitespace
            normalized = re.sub(r'[^\w\s]', '', para.lower())
            normalized = re.sub(r'\s+', ' ', normalized).strip()

            # Skip if we've seen very similar content (same length ± 20% and high overlap)
            is_duplicate = False
            for seen in seen_content:
                # Check length similarity
                if abs(len(normalized) - len(seen)) / max(len(seen), 1) < 0.3:
                    # Check if it's likely the same content in different language
                    # (one is predominantly Latin, one is predominantly Cyrillic)
                    para_latin = sum(1 for c in para if c.isalpha() and ord(c) < 128)
                    para_cyrillic = sum(1 for c in para if '\u0400' <= c <= '\u04FF')
                    total = para_latin + para_cyrillic

                    if total > 0:
                        # If target is English and this paragraph is mostly Latin,
                        # check if we have a Cyrillic version
                        if target_lang.lower() in ['en', 'eng']:
                            # Keep English (Latin) content, skip if duplicate Cyrillic exists
                            if para_latin / total > 0.7:
                                # This is English, keep it
                                pass
                            else:
                                # This is Cyrillic, skip if we already have similar English
                                is_duplicate = True
                                break

            if not is_duplicate:
                seen_content.add(normalized)
                unique_paragraphs.append(para)

        return '\n\n'.join(unique_paragraphs)

    def _looks_like_abbreviation_page(self, text: str) -> bool:
        """
        Detect pages that contain abbreviation/acronym lists.
        Format: "ACRONYM — full expansion" or "ACRONYM - full expansion"
        """
        if not text or len(text) < 50:
            return False

        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        if len(lines) < 3:
            return False

        # Count lines that look like abbreviation definitions
        # Pattern: 2-6 uppercase letters followed by dash/em-dash and explanation
        abbrev_pattern = re.compile(r'^[A-ZА-ЯІЇЄҐ]{2,8}\s*[-—–]\s*.{5,}', re.UNICODE)
        abbrev_lines = sum(1 for ln in lines if abbrev_pattern.match(ln))

        # Also check for header keywords
        header_keywords = ['скорочення', 'позначення', 'абревіатур', 'abbreviation',
                          'acronym', 'умовн', 'condition', 'legend']
        has_header = any(kw in text.lower() for kw in header_keywords)

        # If more than 40% of lines look like abbreviations, or has header + some abbrevs
        return (abbrev_lines >= len(lines) * 0.4) or (has_header and abbrev_lines >= 3)

    async def _translate_abbreviation_page(
        self,
        page_text: str,
        source_lang: str,
        target_lang: str,
        provider: Optional[TranslationProvider]
    ) -> str:
        """
        Translate abbreviation pages while preserving the ACRONYM — expansion format.
        Keeps the original acronym and only translates the expansion.
        CRITICAL: Never allow empty RHS - fallback to original if translation fails.
        """
        out_lines: list[str] = []

        for raw in page_text.splitlines():
            line = raw.strip()
            if not line:
                out_lines.append("")
                continue

            # Pattern: ACRONYM (2-8 chars) followed by dash and explanation
            # Captures: acronym, separator, explanation
            abbrev_match = re.match(
                r'^([A-ZА-ЯІЇЄҐ]{2,8})(\s*[-—–]\s*)(.+)$',
                line,
                re.UNICODE
            )

            if abbrev_match:
                acronym = abbrev_match.group(1)
                explanation = abbrev_match.group(3).strip()

                # Translate only the explanation, keep the acronym
                if self.translation_service.needs_translation(explanation, target_lang):
                    try:
                        translated_exp, _, _ = await self.translation_service.translate_if_needed(
                            explanation, source_lang, target_lang, provider
                        )
                        translated_exp = translated_exp.strip() if translated_exp else ""

                        # CRITICAL: Never allow empty RHS - fallback to original
                        if not translated_exp or len(translated_exp) < 3:
                            logger.warning(f"Empty translation for abbreviation '{acronym}', keeping original")
                            out_lines.append(f"{acronym} — {explanation}")
                        else:
                            out_lines.append(f"{acronym} — {translated_exp}")
                    except Exception as e:
                        # On any error, keep original
                        logger.warning(f"Translation failed for '{acronym}': {e}, keeping original")
                        out_lines.append(f"{acronym} — {explanation}")
                else:
                    # Already in target language
                    out_lines.append(f"{acronym} — {explanation}")
            else:
                # Not an abbreviation line, translate normally
                if self.translation_service.needs_translation(line, target_lang):
                    try:
                        translated_line, _, _ = await self.translation_service.translate_if_needed(
                            line, source_lang, target_lang, provider
                        )
                        translated_line = translated_line.strip() if translated_line else line
                        out_lines.append(translated_line if translated_line else line)
                    except Exception:
                        out_lines.append(line)
                else:
                    out_lines.append(line)

        return "\n".join(out_lines)

    def _split_into_chunks_optimized(self, text: str) -> List[str]:
        """Split into LARGER chunks (5000 chars vs 1000) = 5x faster."""
        chunk_size = self.OPTIMIZED_CHUNK_SIZE
        
        paragraphs = re.split(r'\n\s*\n', text)
        
        chunks = []
        current_chunk = []
        current_size = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            if current_size + len(para) > chunk_size and current_chunk:
                chunks.append('\n\n'.join(current_chunk))
                current_chunk = [para]
                current_size = len(para)
            else:
                current_chunk.append(para)
                current_size += len(para)
        
        if current_chunk:
            chunks.append('\n\n'.join(current_chunk))
        
        return chunks if chunks else ([text] if text else [])

    def _looks_like_toc(self, page_text: str) -> bool:
        """
        Detect TOC-style pages (dot leaders, many short numbered lines).
        Enhanced detection for Ukrainian/Russian documents.
        """
        if not page_text:
            return False

        lines = [ln.strip() for ln in page_text.splitlines() if ln.strip()]
        if len(lines) < 5:
            return False

        # Count various TOC indicators
        dot_leader_lines = sum(1 for ln in lines if "..." in ln or "…" in ln or ". . ." in ln)
        end_number_lines = sum(1 for ln in lines if re.search(r"\b\d{1,4}\s*$", ln))
        short_lines = sum(1 for ln in lines if len(ln) <= 100)

        # Check for TOC header keywords
        text_lower = page_text.lower()
        has_toc_header = any(kw in text_lower for kw in [
            'зміст', 'содержание', 'contents', 'table of contents',
            'розділ', 'раздел', 'chapter', 'глава'
        ])

        # Scoring system for TOC detection
        score = 0

        # Dot leaders are strong indicator
        if dot_leader_lines >= 3:
            score += 2

        # Many lines ending with page numbers
        number_ratio = end_number_lines / max(len(lines), 1)
        if number_ratio >= 0.4:
            score += 2
        elif number_ratio >= 0.25:
            score += 1

        # Mostly short lines (typical for TOC)
        short_ratio = short_lines / max(len(lines), 1)
        if short_ratio >= 0.7:
            score += 1
        elif short_ratio >= 0.5:
            score += 0.5

        # TOC header keyword
        if has_toc_header:
            score += 1

        # Need score >= 2 to be considered TOC
        return score >= 2

    async def _translate_toc_page(
        self,
        page_text: str,
        source_lang: str,
        target_lang: str,
        provider: Optional[TranslationProvider],
    ) -> str:
        """Translate TOC pages line-by-line while preserving page numbers and structure."""
        # Universal document structural terms (language-agnostic, used across all document types)
        # Domain-specific terms (geology, stratigraphy, etc.) should be in the glossary
        structural_replacements = {
            # Table of Contents
            'ЗМІСТ': 'CONTENTS',
            'Зміст': 'Contents',
            'СОДЕРЖАНИЕ': 'CONTENTS',
            'Содержание': 'Contents',
            # Document structure
            'РОЗДІЛ': 'CHAPTER',
            'Розділ': 'Chapter',
            'РАЗДЕЛ': 'CHAPTER',
            'Раздел': 'Chapter',
            'ГЛАВА': 'CHAPTER',
            'Глава': 'Chapter',
            'ЧАСТИНА': 'PART',
            'Частина': 'Part',
            'ЧАСТЬ': 'PART',
            'Часть': 'Part',
            'ПІДРОЗДІЛ': 'SECTION',
            'Підрозділ': 'Section',
            'ПОДРАЗДЕЛ': 'SECTION',
            'Подраздел': 'Section',
            'ПУНКТ': 'ITEM',
            'Пункт': 'Item',
            # Common sections
            'ДОДАТОК': 'APPENDIX',
            'Додаток': 'Appendix',
            'ПРИЛОЖЕНИЕ': 'APPENDIX',
            'Приложение': 'Appendix',
            'ВСТУП': 'INTRODUCTION',
            'Вступ': 'Introduction',
            'ВВЕДЕНИЕ': 'INTRODUCTION',
            'Введение': 'Introduction',
            'ВИСНОВКИ': 'CONCLUSIONS',
            'Висновки': 'Conclusions',
            'ВЫВОДЫ': 'CONCLUSIONS',
            'Выводы': 'Conclusions',
            'ЗАКЛЮЧЕНИЕ': 'CONCLUSION',
            'Заключение': 'Conclusion',
            'ЛІТЕРАТУРА': 'REFERENCES',
            'Література': 'References',
            'ЛИТЕРАТУРА': 'REFERENCES',
            'Литература': 'References',
            'ПЕРЕДМОВА': 'PREFACE',
            'Передмова': 'Preface',
            'ПРЕДИСЛОВИЕ': 'PREFACE',
            'Предисловие': 'Preface',
            'АНОТАЦІЯ': 'ABSTRACT',
            'Анотація': 'Abstract',
            'АННОТАЦИЯ': 'ABSTRACT',
            'Аннотация': 'Abstract',
            'СПИСОК': 'LIST',
            'Список': 'List',
            'ТАБЛИЦЯ': 'TABLE',
            'Таблиця': 'Table',
            'ТАБЛИЦА': 'TABLE',
            'Таблица': 'Table',
            'РИСУНОК': 'FIGURE',
            'Рисунок': 'Figure',
            'МАЛЮНОК': 'FIGURE',
            'Малюнок': 'Figure',
        }

        # Critical proper nouns and stratigraphy terms for TOC
        # These MUST be pre-replaced because LLMs often mistranslate them
        proper_noun_replacements = {
            # Stratigraphy periods (CRITICAL)
            'СТРАТИГРАФІЯ РИФЕЮ': 'STRATIGRAPHY OF THE RIPHEAN',
            'Стратиграфія рифею': 'Stratigraphy of the Riphean',
            'РИФЕЙ': 'RIPHEAN',
            'Рифей': 'Riphean',
            'рифей': 'Riphean',
            'рифейськ': 'Riphean',
            'ВЕНД': 'VENDIAN',
            'Венд': 'Vendian',
            'венд': 'Vendian',
            'ФАНЕРОЗОЙ': 'PHANEROZOIC',
            'Фанерозой': 'Phanerozoic',
            'фанерозой': 'Phanerozoic',
            # Erathem
            'ЕРАТЕМА': 'ERATHEM',
            'Ератема': 'Erathem',
            'ератема': 'erathem',
            # Geographic names (CRITICAL)
            'Ковельськ': 'Kovel',
            'КОВЕЛЬСЬК': 'KOVEL',
            'Переддобрудзьк': 'Fore-Dobruja',
            'ПЕРЕДДОБРУДЗЬК': 'FORE-DOBRUJA',
            'Причорномор': 'Black Sea Region',
            'ПРИЧОРНОМОР': 'BLACK SEA REGION',
            'Західне Причорномор\'я': 'Western Black Sea Region',
            'Волино-Поліськ': 'Volyn-Polissia',
            'ВОЛИНО-ПОЛІСЬК': 'VOLYN-POLISSIA',
            'Овруцьк': 'Ovruch',
            'ОВРУЦЬК': 'OVRUCH',
            'Дніпровсько-Донецьк': 'Dnieper-Donets',
            'ДНІПРОВСЬКО-ДОНЕЦЬК': 'DNIEPER-DONETS',
            # Geological structures
            'прогин': 'Trough',
            'ПРОГИН': 'TROUGH',
            'западина': 'Depression',
            'ЗАПАДИНА': 'DEPRESSION',
            'палеорифт': 'Paleorift',
            'ПАЛЕОРИФТ': 'PALEORIFT',
        }

        out_lines: list[str] = []
        for raw in page_text.splitlines():
            ln = raw.rstrip("\n")
            if not ln.strip():
                out_lines.append(ln)
                continue

            # Apply structural replacements first
            processed_line = ln
            for uk_term, en_term in structural_replacements.items():
                if uk_term in processed_line:
                    processed_line = processed_line.replace(uk_term, en_term)

            # Apply proper noun replacements (critical for accuracy)
            for uk_term, en_term in proper_noun_replacements.items():
                if uk_term in processed_line:
                    processed_line = processed_line.replace(uk_term, en_term)

            # Pattern: left part + dot leader / ellipsis + trailing page number block
            m = re.match(r"^(.*?)(\s*(?:\.{2,}|…)+\s*)(\d{1,4}.*)$", processed_line.strip())
            if m:
                left, leader, right = m.group(1), m.group(2), m.group(3)
                if self.translation_service.needs_translation(left, target_lang):
                    t_left, _, _ = await self.translation_service.translate_if_needed(
                        left, source_lang, target_lang, provider
                    )
                else:
                    t_left = left
                # Reconstruct with proper dot leaders
                out_lines.append(f"{t_left.strip()} {'.' * 20} {right.strip()}")
                continue

            # Handle lines with just trailing numbers (no dot leaders)
            m2 = re.match(r"^(.*?)(\s+)(\d{1,4})\s*$", processed_line.strip())
            if m2:
                left, space, num = m2.group(1), m2.group(2), m2.group(3)
                if self.translation_service.needs_translation(left, target_lang):
                    t_left, _, _ = await self.translation_service.translate_if_needed(
                        left, source_lang, target_lang, provider
                    )
                else:
                    t_left = left
                out_lines.append(f"{t_left.strip()} {'.' * 20} {num}")
                continue

            # Otherwise translate the whole line if needed
            if self.translation_service.needs_translation(ln, target_lang):
                t_ln, _, _ = await self.translation_service.translate_if_needed(
                    ln, source_lang, target_lang, provider
                )
                out_lines.append(t_ln.strip())
            else:
                out_lines.append(ln)

        return "\n".join(out_lines)
    
    def _create_word_document(
        self,
        page_texts: List[str],
        output_path: str,
        total_pages: int,
        source_lang: str,
        target_lang: str,
        pages_translated: int,
        pages_skipped: int,
        page_images: List[List[dict]] = None  # NEW: Images per page
    ):
        """Create Word document from translated pages with images."""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        
        doc = Document()
        
        title = doc.add_heading("Translated Document", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Language: {source_lang} -> {target_lang}").italic = True
        meta.add_run(f"\nTotal Pages: {total_pages}")
        meta.add_run(f"\nPages Translated: {pages_translated} | Pages Kept Original: {pages_skipped}")
        
        # Add image count if available
        if page_images:
            total_images = sum(len(imgs) for imgs in page_images)
            if total_images > 0:
                meta.add_run(f"\nImages Extracted: {total_images}")
        
        doc.add_paragraph()
        doc.add_paragraph("-" * 50)
        doc.add_paragraph()
        
        # Ensure page_images has correct length
        if page_images is None:
            page_images = [[] for _ in range(len(page_texts))]
        while len(page_images) < len(page_texts):
            page_images.append([])
        
        def add_image_to_doc(img_data, img_idx, page_num):
            """Helper to add a single image to the document."""
            try:
                from PIL import Image
                
                img_bytes = img_data["data"]
                img_stream = io.BytesIO(img_bytes)
                
                try:
                    pil_image = Image.open(img_stream)
                    
                    # Convert CMYK or other modes to RGB
                    if pil_image.mode in ('CMYK', 'RGBA', 'LA', 'P'):
                        if pil_image.mode == 'CMYK':
                            pil_image = pil_image.convert('RGB')
                        elif pil_image.mode == 'RGBA':
                            background = Image.new('RGB', pil_image.size, (255, 255, 255))
                            background.paste(pil_image, mask=pil_image.split()[3])
                            pil_image = background
                        elif pil_image.mode == 'P':
                            pil_image = pil_image.convert('RGB')
                        elif pil_image.mode == 'LA':
                            pil_image = pil_image.convert('RGB')
                    
                    output_stream = io.BytesIO()
                    pil_image.save(output_stream, format='PNG')
                    output_stream.seek(0)
                    width = pil_image.width
                    
                except Exception as pil_error:
                    logger.warning(f"PIL processing failed for image {img_idx + 1} on page {page_num}: {pil_error}")
                    output_stream = io.BytesIO(img_bytes)
                    width = img_data.get("width", 400)
                
                # Scale to fit page width (max 6 inches)
                max_width_inches = 6.0
                if width > 0:
                    scale_inches = width / 96.0
                    doc_width = Inches(min(scale_inches, max_width_inches))
                else:
                    doc_width = Inches(4)
                
                doc.add_picture(output_stream, width=doc_width)
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                logger.info(f"Added image {img_idx + 1} to page {page_num} (position: {img_data.get('position', 'unknown')})")
                return True
                
            except Exception as e:
                logger.warning(f"Failed to add image {img_idx + 1} to page {page_num}: {e}")
                p = doc.add_paragraph(f"[Image {img_idx + 1} could not be inserted]")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].italic = True
                return False
        
        for i, page_text in enumerate(page_texts):
            if not page_text.strip() and (i >= len(page_images) or not page_images[i]):
                continue
            
            page_header = doc.add_paragraph()
            page_header.add_run(f"--- Page {i + 1} ---").bold = True
            page_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Get images for this page, separated by position
            current_page_images = page_images[i] if i < len(page_images) else []
            top_images = [img for img in current_page_images if img.get("position") == "top"]
            other_images = [img for img in current_page_images if img.get("position") != "top"]
            
            # Add TOP images first (before text, matching original PDF layout)
            for img_idx, img_data in enumerate(top_images):
                add_image_to_doc(img_data, img_idx, i + 1)
            
            # Add text content
            if page_text.strip():
                paragraphs = page_text.split('\n\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        p = doc.add_paragraph(para_text)
                        p.paragraph_format.space_after = Pt(8)
            
            # Add middle/bottom images after text
            for img_idx, img_data in enumerate(other_images):
                add_image_to_doc(img_data, len(top_images) + img_idx, i + 1)
            
            doc.add_paragraph()
        
        doc.save(output_path)
        logger.info(f"Word document created: {output_path}")
    
    def _generate_output_filename(self, input_path: str, target_lang: str) -> str:
        """Generate output filename as .docx"""
        input_name = Path(input_path).stem
        return f"{input_name}_translated_{target_lang}.docx"
