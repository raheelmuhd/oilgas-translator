"""
Job Processor - FIXED VERSION with Parallel Translation
Handles background document processing with page-level progress tracking.
"""

import asyncio
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Optional, List

import structlog

from app.config import OCRProvider, TranslationProvider, get_settings
from app.models import JobStatus
from app.services.ocr_service import OCRService, OCRProgress
from app.services.translation_service import TranslationService

logger = structlog.get_logger()


class JobProcessor:
    """Processes translation jobs in the background with real-time progress."""
    
    def __init__(self, db_session=None):
        self.settings = get_settings()
        self.ocr_service = OCRService()
        self.translation_service = TranslationService()
        self.db = db_session
    
    async def process_job(
        self,
        job_id: str,
        input_path: str,
        source_lang: Optional[str],
        target_lang: str,
        ocr_provider: Optional[str] = None,
        translation_provider: Optional[str] = None,
        status_callback=None
    ) -> dict:
        """
        Process translation job with smart detection.
        - Extracts text instantly from PDFs with embedded text
        - Only translates non-English content (PARALLEL for speed)
        - Outputs Word document (.docx)
        """
        try:
            # ===== STEP 1: EXTRACT TEXT (0-30%) =====
            if status_callback:
                await status_callback(job_id, JobStatus.EXTRACTING, 0, "Analyzing document...")
            
            loop = asyncio.get_event_loop()
            
            def ocr_progress_wrapper(progress: OCRProgress):
                if status_callback:
                    pct = (progress.current_page / max(progress.total_pages, 1)) * 30
                    asyncio.run_coroutine_threadsafe(
                        status_callback(job_id, JobStatus.EXTRACTING, pct,
                            f"Extracting page {progress.current_page}/{progress.total_pages}"),
                        loop
                    )
            
            ocr_prov = OCRProvider(ocr_provider) if ocr_provider else None
            
            # Get text with page-by-page list
            extracted_text, page_count, ocr_used, page_texts = await self.ocr_service.extract_text(
                input_path, ocr_prov, progress_callback=ocr_progress_wrapper
            )
            
            if not extracted_text.strip():
                raise ValueError("No text could be extracted from the document")
            
            logger.info(f"Extraction complete: {page_count} pages, {len(extracted_text)} chars, method={ocr_used}")
            
            # ===== STEP 2: DETECT LANGUAGE (30-35%) =====
            if status_callback:
                await status_callback(job_id, JobStatus.EXTRACTING, 32, "Detecting language...")
            
            if not source_lang:
                source_lang = await self.translation_service.detect_language(extracted_text)
                logger.info(f"Detected language: {source_lang}")
            
            # ===== STEP 3: SMART TRANSLATION - PARALLEL (35-90%) =====
            if status_callback:
                await status_callback(job_id, JobStatus.TRANSLATING, 35, "Analyzing content for translation...")
            
            trans_prov = TranslationProvider(translation_provider) if translation_provider else None
            
            # First, identify which pages need translation
            pages_needing_translation = []
            pages_to_skip = []
            
            for i, page_text in enumerate(page_texts):
                if not page_text.strip():
                    pages_to_skip.append((i, ""))
                elif not self.translation_service.needs_translation(page_text, target_lang):
                    pages_to_skip.append((i, page_text))  # Already English
                else:
                    pages_needing_translation.append((i, page_text))
            
            pages_skipped = len(pages_to_skip)
            pages_to_translate = len(pages_needing_translation)
            
            logger.info(f"Translation plan: {pages_to_translate} to translate, {pages_skipped} to skip")
            
            if status_callback:
                await status_callback(
                    job_id, JobStatus.TRANSLATING, 38,
                    f"Translating {pages_to_translate} pages ({pages_skipped} already in English)..."
                )
            
            # Translate pages in PARALLEL batches for speed
            BATCH_SIZE = 5  # Translate 5 pages at once
            translated_results = {}
            
            # Add skipped pages to results
            for idx, text in pages_to_skip:
                translated_results[idx] = text
            
            # Translate in parallel batches
            for batch_start in range(0, len(pages_needing_translation), BATCH_SIZE):
                batch = pages_needing_translation[batch_start:batch_start + BATCH_SIZE]
                
                # Create translation tasks for this batch
                async def translate_page(idx: int, text: str) -> tuple:
                    """Translate a single page."""
                    chunks = self._split_into_chunks(text)
                    translated_chunks = []
                    
                    for chunk in chunks:
                        if self.translation_service.needs_translation(chunk, target_lang):
                            result, _, _ = await self.translation_service.translate_if_needed(
                                chunk, source_lang, target_lang, trans_prov
                            )
                            translated_chunks.append(result)
                        else:
                            translated_chunks.append(chunk)
                    
                    return idx, '\n\n'.join(translated_chunks)
                
                # Run batch in parallel
                tasks = [translate_page(idx, text) for idx, text in batch]
                batch_results = await asyncio.gather(*tasks, return_exceptions=True)
                
                # Store results
                for result in batch_results:
                    if isinstance(result, Exception):
                        logger.error(f"Translation error: {result}")
                        continue
                    idx, translated_text = result
                    translated_results[idx] = translated_text
                
                # Update progress
                completed = len([r for r in translated_results.values() if r])
                progress = 38 + (52 * completed / page_count)
                
                if status_callback:
                    await status_callback(
                        job_id, JobStatus.TRANSLATING, progress,
                        f"Translated {min(batch_start + BATCH_SIZE, pages_to_translate)}/{pages_to_translate} pages..."
                    )
            
            # Reconstruct pages in order
            translated_pages = [translated_results.get(i, "") for i in range(page_count)]
            
            logger.info(f"Translation complete: {pages_to_translate} translated, {pages_skipped} skipped")
            
            # ===== STEP 4: CREATE WORD DOCUMENT (90-100%) =====
            if status_callback:
                await status_callback(job_id, JobStatus.TRANSLATING, 92, "Creating Word document...")
            
            output_filename = self._generate_output_filename(input_path, target_lang)
            output_path = Path(self.settings.output_dir) / output_filename
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            self._create_word_document(
                translated_pages, 
                str(output_path), 
                page_count, 
                source_lang, 
                target_lang,
                pages_to_translate,
                pages_skipped
            )
            
            # ===== DONE =====
            if status_callback:
                await status_callback(
                    job_id, JobStatus.COMPLETED, 100, "Translation complete!",
                    {
                        "output_path": str(output_path),
                        "total_pages": page_count,
                        "pages_translated": pages_to_translate,
                        "pages_skipped": pages_skipped
                    }
                )
            
            return {
                "success": True,
                "output_path": str(output_path),
                "page_count": page_count,
                "pages_translated": pages_to_translate,
                "pages_skipped": pages_skipped,
                "source_language": source_lang,
                "target_language": target_lang,
                "ocr_provider": ocr_used
            }
            
        except Exception as e:
            import traceback
            error_msg = str(e)
            logger.error(f"Job {job_id} failed: {error_msg}", traceback=traceback.format_exc())
            if status_callback:
                await status_callback(job_id, JobStatus.FAILED, 0, f"Failed: {error_msg}")
            raise
    
    def _split_into_chunks(self, text: str) -> list:
        """Split text into manageable chunks for translation."""
        chunk_size = self.settings.chunk_size
        
        # Split by paragraphs first
        paragraphs = re.split(r'\n\s*\n', text)
        
        chunks = []
        current_chunk = []
        current_size = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            para_size = len(para)
            
            if current_size + para_size > chunk_size and current_chunk:
                chunks.append('\n\n'.join(current_chunk))
                current_chunk = [para]
                current_size = para_size
            else:
                current_chunk.append(para)
                current_size += para_size
        
        if current_chunk:
            chunks.append('\n\n'.join(current_chunk))
        
        return chunks if chunks else [text]
    
    def _create_word_document(
        self,
        page_texts: List[str],
        output_path: str,
        total_pages: int,
        source_lang: str,
        target_lang: str,
        pages_translated: int,
        pages_skipped: int
    ):
        """Create a professional Word document from translated pages."""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Document title
        title = doc.add_heading("Translated Document", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata paragraph
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Language: {source_lang} → {target_lang}").italic = True
        meta.add_run(f"\nTotal Pages: {total_pages}")
        meta.add_run(f"\nPages Translated: {pages_translated} | Pages Kept Original: {pages_skipped}")
        
        doc.add_paragraph()  # Spacer
        doc.add_paragraph("─" * 50)  # Divider line
        doc.add_paragraph()
        
        # Add each page
        for i, page_text in enumerate(page_texts):
            if not page_text.strip():
                continue
            
            # Page header
            page_header = doc.add_paragraph()
            page_header.add_run(f"━━━ Page {i + 1} ━━━").bold = True
            page_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Page content
            paragraphs = page_text.split('\n\n')
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    p = doc.add_paragraph(para_text)
                    p.paragraph_format.space_after = Pt(8)
            
            doc.add_paragraph()  # Space between pages
        
        # Save the document
        doc.save(output_path)
        logger.info(f"Word document created: {output_path}")
    
    def _generate_output_filename(self, input_path: str, target_lang: str) -> str:
        """Generate output filename as .docx"""
        input_name = Path(input_path).stem
        return f"{input_name}_translated_{target_lang}.docx"
