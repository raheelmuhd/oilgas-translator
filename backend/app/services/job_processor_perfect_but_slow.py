"""
Job Processor - Optimized Sequential Translation Pipeline
==========================================================

FLOW:
1. PARALLEL: Extract all pages from PDF (CPU/OCR)
2. PARALLEL: Pre-chunk all pages (CPU - instant)
3. SEQUENTIAL: Translate page-by-page, chunk-by-chunk (GPU)
   - Each chunk waits for completion before next
   - No queue timeouts (Ollama processes one at a time anyway)
4. Quality validation per chunk
5. Create Word document

This ensures:
- No Ollama queue timeouts
- Each chunk fully translates before moving on
- Proper accuracy tracking
- Retry logic per chunk
"""

import asyncio
import json
import re
import time
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Tuple, Dict

import structlog

from app.config import OCRProvider, TranslationProvider, get_settings
from app.models import JobStatus
from app.services.ocr_service import OCRService, OCRProgress, ImageExtractor
from app.services.translation_service import TranslationService

logger = structlog.get_logger()


# =============================================================================
# CONFIGURATION
# =============================================================================

@dataclass
class TranslationConfig:
    """Pipeline configuration."""
    # Chunking - SMALLER for reliable Ollama processing
    toc_lines_per_chunk: int = 10          # Lines per TOC chunk (reduced)
    narrative_chars_per_chunk: int = 1500  # Chars per chunk (reduced from 2500)
    
    # Timeouts
    chunk_timeout: float = 120.0           # 2 minutes per chunk
    
    # Quality thresholds - lower to avoid infinite retries
    min_acceptable_accuracy: float = 0.30  # Accept 30%+ 
    target_accuracy: float = 0.98          # Ideal accuracy
    
    # Retry - fewer retries
    max_chunk_retries: int = 1             # Only 1 retry
    retry_delay: float = 0.5               # Shorter delay
    
    def log_config(self):
        """Log all configuration values."""
        logger.info("=" * 70)
        logger.info("TRANSLATION CONFIG VALUES:")
        logger.info(f"  narrative_chars_per_chunk = {self.narrative_chars_per_chunk}")
        logger.info(f"  toc_lines_per_chunk       = {self.toc_lines_per_chunk}")
        logger.info(f"  chunk_timeout             = {self.chunk_timeout}s")
        logger.info(f"  min_acceptable_accuracy   = {self.min_acceptable_accuracy:.0%}")
        logger.info(f"  max_chunk_retries         = {self.max_chunk_retries}")
        logger.info(f"  retry_delay               = {self.retry_delay}s")
        logger.info("=" * 70)


# =============================================================================
# METRICS
# =============================================================================

@dataclass
class ChunkResult:
    """Result of translating a single chunk."""
    chunk_index: int
    original_text: str
    translated_text: str
    accuracy: float
    time_taken: float
    attempts: int
    success: bool


@dataclass
class PageResult:
    """Result of translating a full page."""
    page_number: int
    chunks: List[ChunkResult]
    final_text: str
    overall_accuracy: float
    total_time: float
    
    @property
    def success(self) -> bool:
        return self.overall_accuracy >= 0.50


def count_cyrillic(text: str) -> int:
    """Count Cyrillic characters in text."""
    return sum(1 for c in text if '\u0400' <= c <= '\u04FF')


def calculate_accuracy(original: str, translated: str) -> float:
    """Calculate translation accuracy (% of Cyrillic removed)."""
    orig_cyrillic = count_cyrillic(original)
    if orig_cyrillic == 0:
        return 1.0
    remaining = count_cyrillic(translated)
    return max(0.0, 1.0 - (remaining / orig_cyrillic))


def find_hybrid_words(text: str) -> List[str]:
    """Find hybrid words (mixed Latin + Cyrillic)."""
    pattern = r'\b[A-Za-z]+[а-яіїєґА-ЯІЇЄҐ]+\b|\b[а-яіїєґА-ЯІЇЄҐ]+[A-Za-z]+[а-яіїєґА-ЯІЇЄҐ]*\b'
    return list(set(re.findall(pattern, text)))


def check_gpu_available() -> bool:
    """Check if GPU is available."""
    try:
        import torch
        return torch.cuda.is_available()
    except ImportError:
        return False


# =============================================================================
# CHUNKER
# =============================================================================

class SmartChunker:
    """Smart page chunking for optimal translation."""
    
    def __init__(self, config: TranslationConfig):
        self.config = config
        logger.info(f"SmartChunker initialized: max_chars={config.narrative_chars_per_chunk}, toc_lines={config.toc_lines_per_chunk}")
    
    def is_toc(self, text: str) -> bool:
        """Detect if text is table of contents."""
        if not text or len(text) < 200:
            return False
        
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        if len(lines) < 5:
            return False
        
        dot_leaders = sum(1 for ln in lines if '...' in ln or '…' in ln)
        page_numbers = sum(1 for ln in lines if re.search(r'\d{1,4}\s*$', ln))
        
        is_toc = dot_leaders >= 3 or page_numbers >= len(lines) * 0.3
        if is_toc:
            logger.debug(f"TOC detected: {len(lines)} lines, {dot_leaders} dot leaders, {page_numbers} page numbers")
        return is_toc
    
    def chunk_toc(self, text: str) -> List[str]:
        """Chunk TOC by groups of lines."""
        lines = text.split('\n')
        chunks = []
        current = []
        
        for line in lines:
            current.append(line)
            if len(current) >= self.config.toc_lines_per_chunk:
                chunks.append('\n'.join(current))
                current = []
        
        if current:
            chunks.append('\n'.join(current))
        
        logger.debug(f"TOC chunked: {len(lines)} lines → {len(chunks)} chunks of ~{self.config.toc_lines_per_chunk} lines each")
        return chunks if chunks else [text]
    
    def chunk_narrative(self, text: str) -> List[str]:
        """Chunk narrative text by paragraphs."""
        paragraphs = re.split(r'\n\s*\n', text)
        paragraphs = [p.strip() for p in paragraphs if p.strip()]
        
        if not paragraphs:
            return [text] if text.strip() else []
        
        chunks = []
        current = []
        current_size = 0
        max_size = self.config.narrative_chars_per_chunk
        
        for para in paragraphs:
            if current_size + len(para) > max_size and current:
                chunks.append('\n\n'.join(current))
                current = [para]
                current_size = len(para)
            else:
                current.append(para)
                current_size += len(para)
        
        if current:
            chunks.append('\n\n'.join(current))
        
        chunk_sizes = [len(c) for c in chunks]
        logger.debug(f"Narrative chunked: {len(text)} chars → {len(chunks)} chunks, sizes={chunk_sizes}")
        return chunks if chunks else [text]
    
    def chunk_page(self, text: str) -> List[str]:
        """Chunk a page based on content type."""
        if not text.strip():
            logger.debug("Empty page, no chunks")
            return []
        
        if self.is_toc(text):
            chunks = self.chunk_toc(text)
            logger.info(f"Page is TOC: {len(text)} chars → {len(chunks)} chunks")
            return chunks
        else:
            chunks = self.chunk_narrative(text)
            logger.info(f"Page is narrative: {len(text)} chars → {len(chunks)} chunks (max {self.config.narrative_chars_per_chunk} each)")
            return chunks


# =============================================================================
# MAIN JOB PROCESSOR
# =============================================================================

class JobProcessor:
    """
    Optimized translation processor with sequential GPU processing.
    
    Key features:
    - Parallel extraction and chunking (CPU)
    - Sequential translation (GPU - one chunk at a time)
    - Quality validation per chunk
    - Retry logic per chunk
    - No Ollama queue timeouts
    """

    OPTIMIZED_CHUNK_SIZE = 1500  # Updated to match config
    
    def __init__(self, db_session=None):
        self.settings = get_settings()
        self.ocr_service = OCRService()
        self.translation_service = TranslationService()
        self.db = db_session
        self.gpu_available = check_gpu_available()
        
        # Configuration
        self.config = TranslationConfig()
        self.chunker = SmartChunker(self.config)
        
        # Metrics
        self.page_results: Dict[int, PageResult] = {}
        
        # Build safe replacements (complete phrases only!)
        self.safe_replacements = self._build_safe_replacements()
        
        # Log full configuration on startup
        logger.info("")
        logger.info("*" * 70)
        logger.info("*" + " " * 20 + "JobProcessor STARTED" + " " * 28 + "*")
        logger.info("*" * 70)
        logger.info(f"GPU available: {self.gpu_available}")
        logger.info(f"Safe replacements loaded: {len(self.safe_replacements)} terms")
        self.config.log_config()
        logger.info("")

    def _build_safe_replacements(self) -> Dict[str, str]:
        """
        Build SAFE term replacements.
        
        CRITICAL: Only complete phrases! No partial stems like 'прогин' or 'Ковельськ'
        which break Ukrainian declensions and create hybrid words.
        """
        return {
            # Full abbreviation lines (highest priority)
            'ВКМ        — Воронезький кристалічний масив': 'VKM — Voronezh Crystalline Massif',
            'ВКМ — Воронезький кристалічний масив': 'VKM — Voronezh Crystalline Massif',
            'ДДП        — Доно-Дніпровський прогин': 'DDP — Don-Dnieper Trough',
            'ДДП — Доно-Дніпровський прогин': 'DDP — Don-Dnieper Trough',
            'ЛПП        — Львівський палеозойський прогин': 'LPT — Lviv Paleozoic Trough',
            'ЛПП — Львівський палеозойський прогин': 'LPT — Lviv Paleozoic Trough',
            'СЄП       — Східно-Європейська платформа': 'EEP — East European Platform',
            'СЄП — Східно-Європейська платформа': 'EEP — East European Platform',
            'УЩ         — Український щит': 'US — Ukrainian Shield',
            'УЩ — Український щит': 'US — Ukrainian Shield',
            'ДДЗ        — Дніпровсько-Донецька западина': 'DDZ — Dnieper-Donets Depression',
            'ДДЗ — Дніпровсько-Донецька западина': 'DDZ — Dnieper-Donets Depression',
            
            # Complete section headers
            'ПРИЙНЯТІ СКОРОЧЕННЯ': 'ACCEPTED ABBREVIATIONS',
            'Прийняті скорочення': 'Accepted Abbreviations',
            'СТРАТИГРАФІЯ РИФЕЮ': 'STRATIGRAPHY OF THE RIPHEAN',
            'Стратиграфія рифею': 'Stratigraphy of the Riphean',
            'РОЗДІЛ': 'CHAPTER',
            'Розділ': 'Chapter',
            'ВСТУП': 'INTRODUCTION',
            'Вступ': 'Introduction',
            'ЗМІСТ': 'CONTENTS',
            'Зміст': 'Contents',
            'Передмова': 'Preface',
            'ПЕРЕДМОВА': 'PREFACE',
            'Список літератури': 'References',
            'СПИСОК ЛІТЕРАТУРИ': 'REFERENCES',
            
            # Complete system names (full phrase only)
            'ВЕНДСЬКА СИСТЕМА': 'VENDIAN SYSTEM',
            'Вендська система': 'Vendian System',
            'КЕМБРІЙСЬКА СИСТЕМА': 'CAMBRIAN SYSTEM',
            'Кембрійська система': 'Cambrian System',
            'ОРДОВИЦЬКА СИСТЕМА': 'ORDOVICIAN SYSTEM',
            'СИЛУРІЙСЬКА СИСТЕМА': 'SILURIAN SYSTEM',
            'ДЕВОНСЬКА СИСТЕМА': 'DEVONIAN SYSTEM',
            'ТРІАСОВА СИСТЕМА': 'TRIASSIC SYSTEM',
            
            # Complete era names
            'ФАНЕРОЗОЙСЬКА ЕОНОТЕМА': 'PHANEROZOIC EONOTHEM',
            'ПАЛЕОЗОЙСЬКА ЕРАТЕМА': 'PALEOZOIC ERATHEM',
            'МЕЗОЗОЙСЬКА ЕРАТЕМА': 'MESOZOIC ERATHEM',
            
            # Complete location phrases (FULL phrase only)
            'Волино-Поліський прогин': 'Volyn-Polissya Trough',
            'Львівський палеозойський прогин': 'Lviv Paleozoic Trough',
            'Овруцький палеорифт': 'Ovruch Paleorift',
            'Дніпровсько-Донецька западина': 'Dnieper-Donets Depression',
            'Переддобрузький прогин': 'Fore-Dobruja Trough',
            
            # IMPORTANT: The following partial stems are REMOVED to prevent hybrid words:
            # 'прогин' → creates 'Troughу'
            # 'Ковельськ' → creates 'Kovelого'
            # 'ВЕНД' → creates 'VENDIANська'
            # 'рифейськ' → creates 'Ripheanого'
        }

    def _apply_safe_replacements(self, text: str) -> str:
        """Apply safe term replacements (sorted by length, longest first)."""
        if not text:
            return text
        result = text
        for uk, en in sorted(self.safe_replacements.items(), key=lambda x: -len(x[0])):
            if uk in result:
                result = result.replace(uk, en)
        return result

    def _deduplicate_bilingual(self, text: str, target_lang: str) -> str:
        """Remove duplicate bilingual content (e.g., UK abstract + EN abstract)."""
        if not text or len(text) < 100:
            return text

        paragraphs = re.split(r'\n\s*\n', text)
        if len(paragraphs) < 2:
            return text

        latin_paras = []
        cyrillic_paras = []
        other_paras = []
        
        for para in paragraphs:
            para = para.strip()
            if not para or len(para) < 20:
                other_paras.append(para)
                continue
            
            latin = sum(1 for c in para if c.isalpha() and ord(c) < 128)
            cyrillic = count_cyrillic(para)
            total = latin + cyrillic
            
            if total == 0:
                other_paras.append(para)
            elif latin / total > 0.8:
                latin_paras.append(para)
            elif cyrillic / total > 0.8:
                cyrillic_paras.append(para)
            else:
                other_paras.append(para)
        
        # For English target: keep Cyrillic UNLESS there's a matching English version
        if target_lang.lower() in ['en', 'eng']:
            result = list(latin_paras) + list(other_paras)
            
            for cyrillic_para in cyrillic_paras:
                has_english_match = False
                # Only check substantial paragraphs (>300 chars, likely abstracts)
                if len(cyrillic_para) > 300:
                    for latin_para in latin_paras:
                        # Same length ±20% = likely translation pair
                        if abs(len(cyrillic_para) - len(latin_para)) / max(len(latin_para), 1) < 0.2:
                            if len(latin_para) > 300:
                                has_english_match = True
                                break
                
                if not has_english_match:
                    result.append(cyrillic_para)
            
            deduped = '\n\n'.join(result)
            
            # SAFEGUARD: If we removed too much content, keep original
            if len(deduped) < len(text) * 0.3:
                logger.warning(f"Deduplicate removed too much content, keeping original")
                return text
            
            return deduped
        
        return text

    # =========================================================================
    # CORE TRANSLATION LOGIC - SEQUENTIAL CHUNK PROCESSING
    # =========================================================================

    async def _translate_single_chunk(
        self,
        chunk: str,
        chunk_index: int,
        source_lang: str,
        target_lang: str,
        provider: Optional[TranslationProvider]
    ) -> ChunkResult:
        """
        Translate a SINGLE chunk with retry logic.
        
        This is the core sequential operation:
        - Send chunk to GPU/Ollama
        - Wait for completion
        - Validate quality
        - Retry if needed
        - Return result
        
        Only one chunk is processed at a time!
        """
        start_time = time.time()
        
        logger.info(f"    [CHUNK {chunk_index}] Starting translation...")
        logger.info(f"    [CHUNK {chunk_index}] Input: {len(chunk)} chars, {count_cyrillic(chunk)} Cyrillic")
        logger.info(f"    [CHUNK {chunk_index}] First 80 chars: {chunk[:80].replace(chr(10), ' ')}...")
        
        # Empty chunk handling
        if not chunk.strip():
            logger.info(f"    [CHUNK {chunk_index}] Empty chunk, skipping")
            return ChunkResult(
                chunk_index=chunk_index,
                original_text="",
                translated_text="",
                accuracy=1.0,
                time_taken=0.0,
                attempts=0,
                success=True
            )
        
        # Apply safe pre-replacements
        processed_chunk = self._apply_safe_replacements(chunk)
        if processed_chunk != chunk:
            logger.info(f"    [CHUNK {chunk_index}] Safe replacements applied: {len(chunk)} → {len(processed_chunk)} chars")
        
        # Check if translation needed
        if not self.translation_service.needs_translation(processed_chunk, target_lang):
            logger.info(f"    [CHUNK {chunk_index}] No translation needed (already in target language)")
            return ChunkResult(
                chunk_index=chunk_index,
                original_text=chunk,
                translated_text=processed_chunk,
                accuracy=1.0,
                time_taken=0.0,
                attempts=0,
                success=True
            )
        
        # Track best result across retries
        best_translated = processed_chunk
        best_accuracy = 0.0
        attempts = 0
        
        logger.info(f"    [CHUNK {chunk_index}] Calling translator (timeout={self.config.chunk_timeout}s, max_retries={self.config.max_chunk_retries})...")
        
        for attempt in range(self.config.max_chunk_retries + 1):
            attempts = attempt + 1
            logger.info(f"    [CHUNK {chunk_index}] Attempt {attempts}/{self.config.max_chunk_retries + 1}...")
            
            try:
                # Translate with timeout - SEQUENTIAL (wait for result)
                call_start = time.time()
                translated, _, _ = await asyncio.wait_for(
                    self.translation_service.translate_if_needed(
                        processed_chunk, source_lang, target_lang, provider
                    ),
                    timeout=self.config.chunk_timeout
                )
                call_duration = time.time() - call_start
                
                logger.info(f"    [CHUNK {chunk_index}] Translator returned in {call_duration:.1f}s")
                logger.info(f"    [CHUNK {chunk_index}] Output: {len(translated) if translated else 0} chars")
                
                # Detect complete failure cases
                if not translated or not translated.strip():
                    logger.warning(f"    [CHUNK {chunk_index}] ⚠️ EMPTY RESPONSE from translator!")
                    if attempt < self.config.max_chunk_retries:
                        logger.info(f"    [CHUNK {chunk_index}] Waiting {self.config.retry_delay}s before retry...")
                        await asyncio.sleep(self.config.retry_delay)
                    continue
                
                # Check if returned text is identical to input (no translation)
                if translated.strip() == processed_chunk.strip():
                    logger.warning(f"    [CHUNK {chunk_index}] ⚠️ TRANSLATOR RETURNED ORIGINAL TEXT UNCHANGED!")
                    logger.warning(f"    [CHUNK {chunk_index}] This means Ollama did NOT translate. Breaking retry loop.")
                    # Don't retry - translator isn't working for this chunk
                    best_translated = processed_chunk
                    best_accuracy = 0.0
                    break
                
                # Calculate accuracy
                accuracy = calculate_accuracy(chunk, translated)
                remaining_cyrillic = count_cyrillic(translated)
                
                logger.info(f"    [CHUNK {chunk_index}] Accuracy: {accuracy:.0%} (Cyrillic: {count_cyrillic(chunk)} → {remaining_cyrillic})")
                
                # Log what we got
                if accuracy == 0.0:
                    logger.warning(
                        f"    [CHUNK {chunk_index}] ⚠️ 0% ACCURACY - response still has all Cyrillic!"
                    )
                    logger.warning(f"    [CHUNK {chunk_index}] Response first 150 chars: {translated[:150].replace(chr(10), ' ')}...")
                elif accuracy < self.config.min_acceptable_accuracy:
                    logger.warning(f"    [CHUNK {chunk_index}] ⚠️ LOW ACCURACY: {accuracy:.0%} < {self.config.min_acceptable_accuracy:.0%} threshold")
                
                # Track best result
                if accuracy > best_accuracy:
                    best_translated = translated
                    best_accuracy = accuracy
                    logger.info(f"    [CHUNK {chunk_index}] New best accuracy: {best_accuracy:.0%}")
                
                # Check if good enough
                if accuracy >= self.config.min_acceptable_accuracy:
                    elapsed = time.time() - start_time
                    logger.info(f"    [CHUNK {chunk_index}] ✓ SUCCESS: {accuracy:.0%} >= {self.config.min_acceptable_accuracy:.0%} threshold")
                    logger.info(f"    [CHUNK {chunk_index}] Total time: {elapsed:.1f}s, attempts: {attempts}")
                    
                    return ChunkResult(
                        chunk_index=chunk_index,
                        original_text=chunk,
                        translated_text=translated,
                        accuracy=accuracy,
                        time_taken=elapsed,
                        attempts=attempts,
                        success=True
                    )
                
                # Low accuracy - retry if attempts remain
                if attempt < self.config.max_chunk_retries:
                    logger.warning(
                        f"    [CHUNK {chunk_index}] Accuracy {accuracy:.0%} below threshold, retrying..."
                    )
                    await asyncio.sleep(self.config.retry_delay)
                
            except asyncio.TimeoutError:
                logger.error(f"    [CHUNK {chunk_index}] ⚠️ TIMEOUT after {self.config.chunk_timeout}s (attempt {attempts})")
                if attempt < self.config.max_chunk_retries:
                    logger.info(f"    [CHUNK {chunk_index}] Will retry after {self.config.retry_delay}s...")
                    await asyncio.sleep(self.config.retry_delay)
                    
            except Exception as e:
                logger.error(f"    [CHUNK {chunk_index}] ⚠️ ERROR: {e}")
                import traceback
                logger.error(f"    [CHUNK {chunk_index}] Traceback: {traceback.format_exc()}")
                if attempt < self.config.max_chunk_retries:
                    await asyncio.sleep(self.config.retry_delay)
        
        # Return best result after all retries
        elapsed = time.time() - start_time
        success = best_accuracy >= self.config.min_acceptable_accuracy
        
        logger.warning(f"    [CHUNK {chunk_index}] === FINAL RESULT ===")
        logger.warning(f"    [CHUNK {chunk_index}] Accuracy: {best_accuracy:.0%}")
        logger.warning(f"    [CHUNK {chunk_index}] Attempts: {attempts}")
        logger.warning(f"    [CHUNK {chunk_index}] Time: {elapsed:.1f}s")
        logger.warning(f"    [CHUNK {chunk_index}] Success: {success}")
        
        return ChunkResult(
            chunk_index=chunk_index,
            original_text=chunk,
            translated_text=best_translated,
            accuracy=best_accuracy,
            time_taken=elapsed,
            attempts=attempts,
            success=success
        )

    async def _translate_page_sequential(
        self,
        page_text: str,
        page_number: int,
        source_lang: str,
        target_lang: str,
        provider: Optional[TranslationProvider]
    ) -> PageResult:
        """
        Translate a full page by processing chunks SEQUENTIALLY.
        
        Flow:
        1. Pre-process (deduplicate, apply replacements)
        2. Chunk the page
        3. For each chunk (SEQUENTIAL):
           - Translate chunk → WAIT for result
           - Validate accuracy
           - Move to next chunk only after completion
        4. Concatenate all translated chunks
        5. Return page result with metrics
        """
        start_time = time.time()
        
        logger.info("=" * 50)
        logger.info(f"PAGE {page_number}: Starting translation")
        logger.info(f"  Original text length: {len(page_text)} chars")
        logger.info(f"  Cyrillic chars in original: {count_cyrillic(page_text)}")
        
        # Handle empty pages
        if not page_text.strip():
            logger.info(f"PAGE {page_number}: Empty page, skipping")
            return PageResult(
                page_number=page_number,
                chunks=[],
                final_text="",
                overall_accuracy=1.0,
                total_time=0.0
            )
        
        # Pre-process: deduplicate bilingual content
        original_len = len(page_text)
        processed_text = self._deduplicate_bilingual(page_text, target_lang)
        if len(processed_text) != original_len:
            logger.info(f"PAGE {page_number}: Deduplicated {original_len} → {len(processed_text)} chars")
        
        # Chunk the page (CPU operation, instant)
        chunks = self.chunker.chunk_page(processed_text)
        
        if not chunks:
            logger.warning(f"PAGE {page_number}: No chunks created from text")
            return PageResult(
                page_number=page_number,
                chunks=[],
                final_text=processed_text,
                overall_accuracy=1.0,
                total_time=0.0
            )
        
        # Log chunk details
        logger.info(f"PAGE {page_number}: Created {len(chunks)} chunks:")
        for i, chunk in enumerate(chunks):
            cyrillic_count = count_cyrillic(chunk)
            logger.info(f"  Chunk {i}: {len(chunk)} chars, {cyrillic_count} Cyrillic")
        
        # Translate chunks ONE BY ONE (sequential!)
        chunk_results: List[ChunkResult] = []
        translated_texts: List[str] = []
        
        for i, chunk in enumerate(chunks):
            logger.info(f"PAGE {page_number} CHUNK {i+1}/{len(chunks)}: Starting translation...")
            logger.info(f"  Input: {len(chunk)} chars, {count_cyrillic(chunk)} Cyrillic")
            
            # SEQUENTIAL: Wait for each chunk to complete before starting next
            result = await self._translate_single_chunk(
                chunk, i, source_lang, target_lang, provider
            )
            
            chunk_results.append(result)
            translated_texts.append(result.translated_text)
            
            # Log detailed chunk result
            logger.info(
                f"PAGE {page_number} CHUNK {i+1}/{len(chunks)}: DONE - "
                f"{result.accuracy:.0%} accuracy, {result.time_taken:.1f}s, "
                f"attempts={result.attempts}, success={result.success}"
            )
            logger.info(f"  Output: {len(result.translated_text)} chars, {count_cyrillic(result.translated_text)} Cyrillic remaining")
        
        # Concatenate all translated chunks
        final_text = '\n\n'.join(translated_texts)
        
        # Calculate overall page metrics
        total_orig_cyrillic = sum(count_cyrillic(r.original_text) for r in chunk_results)
        total_remaining = sum(count_cyrillic(r.translated_text) for r in chunk_results)
        
        if total_orig_cyrillic > 0:
            overall_accuracy = 1.0 - (total_remaining / total_orig_cyrillic)
        else:
            overall_accuracy = 1.0
        
        total_time = time.time() - start_time
        
        page_result = PageResult(
            page_number=page_number,
            chunks=chunk_results,
            final_text=final_text,
            overall_accuracy=overall_accuracy,
            total_time=total_time
        )
        
        # Log page completion summary
        successful_chunks = sum(1 for r in chunk_results if r.success)
        logger.info("=" * 50)
        logger.info(f"PAGE {page_number}: COMPLETE")
        logger.info(f"  Overall accuracy: {overall_accuracy:.0%}")
        logger.info(f"  Chunks: {successful_chunks}/{len(chunks)} successful")
        logger.info(f"  Cyrillic: {total_orig_cyrillic} → {total_remaining} ({total_remaining} remaining)")
        logger.info(f"  Time: {total_time:.1f}s")
        logger.info("=" * 50)
        
        return page_result

    # =========================================================================
    # PARALLEL PRE-CHUNKING (CPU-bound, fast)
    # =========================================================================

    def _prepare_all_chunks(self, page_texts: List[str]) -> List[List[str]]:
        """
        Pre-chunk ALL pages in parallel (CPU operation, instant).
        
        Returns: List of chunk lists, one per page
        """
        all_chunks = []
        total_chunk_count = 0
        
        for page_text in page_texts:
            if page_text.strip():
                chunks = self.chunker.chunk_page(page_text)
                total_chunk_count += len(chunks)
            else:
                chunks = []
            all_chunks.append(chunks)
        
        logger.info(
            f"Pre-chunked {len(page_texts)} pages into {total_chunk_count} total chunks"
        )
        
        return all_chunks

    # =========================================================================
    # DEBUG HELPERS
    # =========================================================================

    def _debug_dir(self, job_id: str) -> Path:
        settings = get_settings()
        base = Path(settings.debug_output_dir)
        d = base / str(job_id)
        d.mkdir(parents=True, exist_ok=True)
        return d

    def _debug_write_text(self, job_id: str, filename: str, content: str):
        settings = get_settings()
        if not getattr(settings, "debug_translation", False):
            return
        d = self._debug_dir(job_id)
        max_chars = int(getattr(settings, "debug_max_chars_per_page", 200000))
        safe = content if content is not None else ""
        if max_chars > 0 and len(safe) > max_chars:
            safe = safe[:max_chars] + "\n\n[TRUNCATED]\n"
        (d / filename).write_text(safe, encoding="utf-8", errors="ignore")

    def _debug_write_json(self, job_id: str, filename: str, data: dict):
        settings = get_settings()
        if not getattr(settings, "debug_translation", False):
            return
        d = self._debug_dir(job_id)
        (d / filename).write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8", errors="ignore")

    # =========================================================================
    # QUALITY REPORT
    # =========================================================================

    def _generate_quality_report(self) -> str:
        """Generate a quality report from collected metrics."""
        if not self.page_results:
            return "No metrics collected"
        
        total_pages = len(self.page_results)
        total_chunks = sum(len(r.chunks) for r in self.page_results.values())
        successful_chunks = sum(
            sum(1 for c in r.chunks if c.success) 
            for r in self.page_results.values()
        )
        
        avg_accuracy = sum(r.overall_accuracy for r in self.page_results.values()) / total_pages
        total_time = sum(r.total_time for r in self.page_results.values())
        
        failed_pages = [p for p, r in self.page_results.items() if r.overall_accuracy < 0.5]
        low_quality_pages = [p for p, r in self.page_results.items() if 0.5 <= r.overall_accuracy < 0.9]
        
        report = f"""
============================================================
TRANSLATION QUALITY REPORT
============================================================
Total pages translated: {total_pages}
Total chunks processed: {total_chunks}
Successful chunks: {successful_chunks}/{total_chunks} ({successful_chunks/max(total_chunks,1)*100:.1f}%)

Average accuracy: {avg_accuracy:.1%}
Total time: {total_time:.1f}s ({total_time/max(total_pages,1):.1f}s per page)

Failed pages (<50% accuracy): {failed_pages if failed_pages else 'None'}
Low quality pages (50-90%): {low_quality_pages if low_quality_pages else 'None'}
============================================================
"""
        return report

    # =========================================================================
    # WORD DOCUMENT CREATION
    # =========================================================================

    def _create_word_document(
        self,
        page_texts: List[str],
        output_path: str,
        total_pages: int,
        source_lang: str,
        target_lang: str,
        pages_translated: int,
        pages_skipped: int,
        page_images: List[List[dict]] = None
    ):
        """Create Word document from translated pages with images."""
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        
        doc = Document()
        
        title = doc.add_heading("Translated Document", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Language: {source_lang} → {target_lang}").italic = True
        meta.add_run(f"\nTotal Pages: {total_pages}")
        meta.add_run(f"\nPages Translated: {pages_translated} | Skipped: {pages_skipped}")
        meta.add_run(f"\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        if page_images:
            total_images = sum(len(imgs) for imgs in page_images)
            if total_images > 0:
                meta.add_run(f"\nImages Extracted: {total_images}")
        
        doc.add_paragraph()
        doc.add_paragraph("-" * 50)
        doc.add_paragraph()
        
        # Ensure page_images has correct length
        if page_images is None:
            page_images = [[] for _ in range(len(page_texts))]
        while len(page_images) < len(page_texts):
            page_images.append([])
        
        def add_image_to_doc(img_data, img_idx, page_num):
            """Helper to add a single image to the document."""
            try:
                from PIL import Image
                
                img_bytes = img_data.get("data") or img_data.get("bytes")
                if not img_bytes:
                    return False
                    
                img_stream = io.BytesIO(img_bytes)
                
                try:
                    pil_image = Image.open(img_stream)
                    
                    # Convert modes
                    if pil_image.mode in ('CMYK', 'RGBA', 'LA', 'P'):
                        if pil_image.mode == 'CMYK':
                            pil_image = pil_image.convert('RGB')
                        elif pil_image.mode == 'RGBA':
                            background = Image.new('RGB', pil_image.size, (255, 255, 255))
                            background.paste(pil_image, mask=pil_image.split()[3])
                            pil_image = background
                        elif pil_image.mode in ('P', 'LA'):
                            pil_image = pil_image.convert('RGB')
                    
                    output_stream = io.BytesIO()
                    pil_image.save(output_stream, format='PNG')
                    output_stream.seek(0)
                    width = pil_image.width
                    
                except Exception as pil_error:
                    logger.warning(f"PIL processing failed: {pil_error}")
                    output_stream = io.BytesIO(img_bytes)
                    width = img_data.get("width", 400)
                
                max_width_inches = 6.0
                if width > 0:
                    scale_inches = width / 96.0
                    doc_width = Inches(min(scale_inches, max_width_inches))
                else:
                    doc_width = Inches(4)
                
                doc.add_picture(output_stream, width=doc_width)
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                return True
                
            except Exception as e:
                logger.warning(f"Failed to add image {img_idx + 1}: {e}")
                return False
        
        for i, page_text in enumerate(page_texts):
            # ALWAYS add page marker (ensures no missing pages)
            page_header = doc.add_paragraph()
            page_header.add_run(f"--- Page {i + 1} ---").bold = True
            page_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add quality warning if page had issues
            if (i + 1) in self.page_results:
                result = self.page_results[i + 1]
                if result.overall_accuracy < 0.5:
                    warning = doc.add_paragraph()
                    run = warning.add_run(
                        f"⚠️ Low translation quality: {result.overall_accuracy:.0%}"
                    )
                    run.italic = True
                    run.font.color.rgb = RGBColor(255, 140, 0)
            
            # Handle empty pages
            if not page_text.strip() and (i >= len(page_images) or not page_images[i]):
                empty_note = doc.add_paragraph()
                empty_note.add_run("[Empty page in original]").italic = True
                doc.add_paragraph()
                continue
            
            # Get images for this page
            current_page_images = page_images[i] if i < len(page_images) else []
            top_images = [img for img in current_page_images if img.get("position") == "top"]
            other_images = [img for img in current_page_images if img.get("position") != "top"]
            
            # Add TOP images first
            for img_idx, img_data in enumerate(top_images):
                add_image_to_doc(img_data, img_idx, i + 1)
            
            # Add text content
            if page_text.strip():
                paragraphs = page_text.split('\n\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        p = doc.add_paragraph(para_text)
                        p.paragraph_format.space_after = Pt(8)
            
            # Add other images after text
            for img_idx, img_data in enumerate(other_images):
                add_image_to_doc(img_data, len(top_images) + img_idx, i + 1)
            
            doc.add_paragraph()
        
        doc.save(output_path)
        logger.info(f"Word document created: {output_path}")
    
    def _generate_output_filename(self, input_path: str, target_lang: str) -> str:
        """Generate output filename."""
        input_name = Path(input_path).stem
        return f"{input_name}_translated_{target_lang}.docx"

    # =========================================================================
    # MAIN PROCESS JOB
    # =========================================================================

    async def process_job(
        self,
        job_id: str,
        input_path: str,
        source_lang: Optional[str],
        target_lang: str,
        ocr_provider: Optional[str] = None,
        translation_provider: Optional[str] = None,
        status_callback=None
    ) -> dict:
        """
        Process translation job with optimized pipeline:
        
        1. PARALLEL: Extract all pages from PDF (via ocr_service)
        2. PARALLEL: Pre-chunk all pages (CPU, instant)
        3. SEQUENTIAL: Translate page-by-page, chunk-by-chunk (GPU)
        4. Create Word document with quality markers
        """
        try:
            # Reset metrics for new job
            self.page_results = {}
            
            # Determine provider
            provider_name = translation_provider or "ollama"
            ollama_available = False
            ollama_model = self.settings.ollama_model
            
            try:
                ollama_translator = self.translation_service.providers.get(TranslationProvider.OLLAMA)
                if ollama_translator:
                    ollama_available = ollama_translator.is_available()
                    logger.info(f"Ollama availability: {ollama_available}")
            except Exception as e:
                logger.error(f"Error checking Ollama: {e}")

            if not translation_provider:
                if self.gpu_available and ollama_available:
                    provider_name = "ollama"
                    logger.info(f"Using Ollama/{ollama_model} with GPU")
                elif ollama_available:
                    provider_name = "ollama"
                    logger.info(f"Using Ollama/{ollama_model} on CPU")
                else:
                    provider_name = "nllb"

            device_info = f"{ollama_model} ({'GPU' if self.gpu_available else 'CPU'})"
            
            # ===== STEP 1: EXTRACT TEXT (PARALLEL via ocr_service) =====
            if status_callback:
                await status_callback(job_id, JobStatus.EXTRACTING, 0, "Analyzing document...")
            
            loop = asyncio.get_event_loop()
            
            def ocr_progress_wrapper(progress: OCRProgress):
                if status_callback:
                    pct = (progress.current_page / max(progress.total_pages, 1)) * 18
                    asyncio.run_coroutine_threadsafe(
                        status_callback(job_id, JobStatus.EXTRACTING, pct,
                            f"Extracting page {progress.current_page}/{progress.total_pages}"),
                        loop
                    )
            
            ocr_prov = OCRProvider(ocr_provider) if ocr_provider else None
            
            extracted_text, page_count, ocr_used, page_texts = await self.ocr_service.extract_text(
                input_path, ocr_prov, progress_callback=ocr_progress_wrapper
            )
            
            # Debug output
            settings = get_settings()
            if getattr(settings, "debug_translation", False):
                self._debug_write_text(job_id, "extracted_full.txt", extracted_text or "")
                for i, page_text in enumerate(page_texts or []):
                    self._debug_write_text(job_id, f"page_{i:03d}_extracted.txt", page_text or "")
            
            if not extracted_text.strip():
                raise ValueError("No text could be extracted from the document")
            
            logger.info(f"Extraction complete: {page_count} pages, {len(extracted_text)} chars")
            
            # ===== STEP 1B: EXTRACT IMAGES =====
            if status_callback:
                await status_callback(job_id, JobStatus.EXTRACTING, 18, "Extracting images...")
            
            page_images = []
            try:
                page_images = await ImageExtractor.extract_images_from_pdf(input_path)
                total_images = sum(len(imgs) for imgs in page_images)
                logger.info(f"Image extraction: {total_images} images from {len(page_images)} pages")
            except Exception as e:
                logger.warning(f"Image extraction failed: {e}")
                page_images = [[] for _ in range(page_count)]
            
            # ===== STEP 2: DETECT LANGUAGE =====
            if status_callback:
                await status_callback(job_id, JobStatus.EXTRACTING, 20, "Detecting language...")
            
            if not source_lang:
                source_lang = await self.translation_service.detect_language(extracted_text)
                logger.info(f"Detected language: {source_lang}")
            
            # ===== STEP 3: PRE-CHUNK ALL PAGES (PARALLEL - CPU) =====
            if status_callback:
                await status_callback(job_id, JobStatus.TRANSLATING, 22, "Preparing chunks...")
            
            # This is CPU-bound and instant
            all_page_chunks = self._prepare_all_chunks(page_texts)
            
            # Separate pages needing translation
            pages_to_translate = []
            pages_to_skip = []
            
            for i, page_text in enumerate(page_texts):
                if not page_text.strip():
                    pages_to_skip.append((i, ""))
                elif not self.translation_service.needs_translation(page_text, target_lang):
                    pages_to_skip.append((i, page_text))
                else:
                    pages_to_translate.append((i, page_text))
            
            pages_skipped = len(pages_to_skip)
            total_to_translate = len(pages_to_translate)
            
            logger.info(f"To translate: {total_to_translate} pages, skip: {pages_skipped}")
            
            speed_note = f"Using {provider_name} on {device_info}"
            if status_callback:
                await status_callback(
                    job_id, JobStatus.TRANSLATING, 25,
                    f"Translating {total_to_translate} pages sequentially... ({speed_note})",
                    {"pages_skipped": pages_skipped, "total_pages": page_count}
                )
            
            # Initialize results with skipped pages
            translated_results = {idx: text for idx, text in pages_to_skip}
            
            trans_prov = TranslationProvider(translation_provider) if translation_provider else None
            
            # ===== STEP 4: TRANSLATE PAGES SEQUENTIALLY =====
            logger.info("=" * 70)
            logger.info("STARTING SEQUENTIAL PAGE TRANSLATION")
            logger.info(f"  Total pages to translate: {total_to_translate}")
            logger.info(f"  Chunk size: {self.config.narrative_chars_per_chunk} chars")
            logger.info(f"  Timeout per chunk: {self.config.chunk_timeout}s")
            logger.info(f"  Max retries: {self.config.max_chunk_retries}")
            logger.info(f"  Min acceptable accuracy: {self.config.min_acceptable_accuracy:.0%}")
            logger.info("=" * 70)
            
            for i, (idx, page_text) in enumerate(pages_to_translate):
                logger.info(f"\n>>> TRANSLATING PAGE {i+1}/{total_to_translate} (page index {idx+1}) <<<\n")
                
                # Translate this page (chunks processed sequentially within)
                page_result = await self._translate_page_sequential(
                    page_text, idx + 1, source_lang, target_lang, trans_prov
                )
                
                translated_results[idx] = page_result.final_text
                self.page_results[idx + 1] = page_result
                
                # Log progress summary
                elapsed_pages = i + 1
                remaining_pages = total_to_translate - elapsed_pages
                avg_time = page_result.total_time
                eta = remaining_pages * avg_time
                
                logger.info(f">>> PAGE {i+1}/{total_to_translate} DONE: {page_result.overall_accuracy:.0%} in {page_result.total_time:.1f}s <<<")
                logger.info(f"    Remaining: {remaining_pages} pages, ETA: ~{eta:.0f}s")
                
                # Progress update
                progress = 25 + (65 * (i + 1) / max(total_to_translate, 1))
                
                if status_callback:
                    await status_callback(
                        job_id, JobStatus.TRANSLATING, min(progress, 90),
                        f"Page {i + 1}/{total_to_translate}: {page_result.overall_accuracy:.0%} accuracy",
                        {"page": idx + 1, "accuracy": page_result.overall_accuracy}
                    )
            
            logger.info("=" * 70)
            logger.info("TRANSLATION PHASE COMPLETE")
            logger.info("=" * 70)
            
            # Reconstruct in order
            translated_pages = [translated_results.get(i, "") for i in range(page_count)]
            
            # Debug output
            if getattr(settings, "debug_translation", False):
                for i, t in enumerate(translated_pages or []):
                    self._debug_write_text(job_id, f"page_{i:03d}_translated.txt", t or "")
                self._debug_write_text(job_id, "translated_full.txt", "\n\n".join(translated_pages or []))
            
            # Generate quality report
            quality_report = self._generate_quality_report()
            logger.info(quality_report)
            
            logger.info(f"Translation complete: {total_to_translate} translated, {pages_skipped} skipped")
            
            # ===== STEP 5: CREATE WORD DOCUMENT =====
            if status_callback:
                await status_callback(job_id, JobStatus.TRANSLATING, 92, "Creating Word document...")
            
            output_filename = self._generate_output_filename(input_path, target_lang)
            output_path = Path(self.settings.output_dir) / output_filename
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            self._create_word_document(
                translated_pages, 
                str(output_path), 
                page_count, 
                source_lang, 
                target_lang,
                total_to_translate,
                pages_skipped,
                page_images
            )
            
            # ===== DONE =====
            failed_pages = [p for p, r in self.page_results.items() if r.overall_accuracy < 0.5]
            
            completion_msg = "Translation complete!"
            if failed_pages:
                completion_msg = f"Complete! ({len(failed_pages)} pages may need review)"
            
            if status_callback:
                await status_callback(
                    job_id, JobStatus.COMPLETED, 100, completion_msg,
                    {
                        "output_path": str(output_path),
                        "total_pages": page_count,
                        "pages_translated": total_to_translate,
                        "pages_skipped": pages_skipped,
                        "failed_pages": failed_pages,
                        "quality_report": quality_report
                    }
                )
            
            return {
                "success": True,
                "output_path": str(output_path),
                "page_count": page_count,
                "pages_translated": total_to_translate,
                "pages_skipped": pages_skipped,
                "failed_pages": failed_pages,
                "source_language": source_lang,
                "target_language": target_lang,
                "ocr_provider": ocr_used,
                "translation_provider": provider_name,
                "gpu_used": self.gpu_available,
                "quality_report": quality_report
            }
            
        except Exception as e:
            import traceback
            error_msg = str(e)
            logger.error(f"Job {job_id} failed: {error_msg}", traceback=traceback.format_exc())
            if status_callback:
                await status_callback(job_id, JobStatus.FAILED, 0, f"Failed: {error_msg}")
            raise
